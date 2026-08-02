# -*- coding: utf-8 -*-
"""
Informe de Estado de Tareas Trello - Generico multi-cliente
=============================================================
Uso:  python generar_informe.py [carpeta_destino]

Si no se pasa carpeta_destino, usa la carpeta del script.
Busca el JSON de Trello mas nuevo en esa carpeta y genera:
  - tareas_sgc_YYYY-MM-DD.xlsx / .csv     (Parte 1)
  - informe_interno_YYYY-MM-DD.html       (Parte 2)
  - informe_direccion_YYYY-MM-DD.docx     (Parte 3)

Config por cliente (config.json):
  Se busca primero en carpeta_destino, despues en su carpeta padre
  (la raiz _trello). Debe tener:
    { "nombre_laboratorio": "...", "normas": ["9001","17025",...] }
  Si no se encuentra, usa valores por defecto genericos.

Si existe un archivo 'observaciones.json' en la carpeta destino con las
claves novedades / decisiones / compromisos (listas de strings), esos
textos se usan para completar la Seccion 5 del Word en lugar de dejarla
en blanco.

Requisitos:
  pip install pandas openpyxl matplotlib python-docx
"""
import json, os, glob, sys, io, base64
from datetime import datetime, timedelta
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Config
SCRIPT_DIR  = sys.argv[1] if len(sys.argv) > 1 else os.path.dirname(os.path.abspath(__file__))
SCRIPT_DIR  = os.path.abspath(SCRIPT_DIR)

ARCHIVOS_NO_TRELLO = ('observaciones.json', 'estado.json', 'config.json')

def cargar_config(carpeta):
    for candidato in (os.path.join(carpeta, 'config.json'),
                       os.path.join(os.path.dirname(carpeta), 'config.json')):
        if os.path.exists(candidato):
            with open(candidato, 'r', encoding='utf-8') as f:
                return json.load(f)
    return {}

_cfg = cargar_config(SCRIPT_DIR)
LAB_NAME    = _cfg.get('nombre_laboratorio', 'Laboratorio')
NORMAS      = _cfg.get('normas', ['9001'])
PRIORIDADES = ['Prioridad alta', 'Prioridad media', 'Prioridad baja']
MESES_ES    = {'January':'enero','February':'febrero','March':'marzo',
               'April':'abril','May':'mayo','June':'junio','July':'julio',
               'August':'agosto','September':'septiembre','October':'octubre',
               'November':'noviembre','December':'diciembre'}

C_AZUL_OSC='1F3864'; C_AZUL_MED='2E75B6'; C_BLANCO='FFFFFF'
C_VERDE='00703A';    C_VERDE_CLR='E8F5E9'
C_ROJO='C00000';     C_ROJO_CLR='FFEEEE'
C_NARANJA='FF9900';  C_NARJ_CLR='FFF8E7'

json_files = [f for f in glob.glob(os.path.join(SCRIPT_DIR, '*.json'))
              if os.path.basename(f) not in ARCHIVOS_NO_TRELLO]
if not json_files:
    sys.exit("ERROR: No se encontro ningun archivo JSON de Trello en " + SCRIPT_DIR)
JSON_PATH = sorted(json_files)[-1]

OBS_PATH = os.path.join(SCRIPT_DIR, 'observaciones.json')
observaciones = None
if os.path.exists(OBS_PATH):
    with open(OBS_PATH, 'r', encoding='utf-8') as f:
        observaciones = json.load(f)

hoy        = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
fecha_str  = datetime.now().strftime('%Y-%m-%d')
semana_num = hoy.isocalendar()[1]
fecha_larga = datetime.now().strftime('%d de %B de %Y')
for en, es in MESES_ES.items():
    fecha_larga = fecha_larga.replace(en, es)

SEP = '=' * 60
print('\n' + SEP)
print('  Generador de Informes Trello - ' + LAB_NAME)
print('  Carpeta: ' + SCRIPT_DIR)
print('  Archivo: ' + os.path.basename(JSON_PATH))
print('  Normas : ' + ', '.join(NORMAS))
print('  Observaciones IA: ' + ('SI' if observaciones else 'NO (plantilla en blanco)'))
print('  Fecha  : ' + datetime.now().strftime('%d/%m/%Y %H:%M'))
print(SEP + '\n')

# PARTE 1 - Procesar JSON -> CSV / Excel
print("[1/3] Procesando tablero Trello...")

with open(JSON_PATH, 'r', encoding='utf-8') as f:
    trello_data = json.load(f)

list_dict   = {l['id']: l['name'] for l in trello_data.get('lists', [])}
member_dict = {m['id']: m.get('fullName', m.get('username','?'))
               for m in trello_data.get('members', [])}

task_data = []
for card in trello_data.get('cards', []):
    if card.get('closed'): continue
    nombre = card.get('name','').strip()
    if 'Categoria' in nombre or '[Categoria]' in nombre: continue

    labels  = [l.get('name') for l in card.get('labels',[]) if l.get('name')]
    members = [member_dict.get(mid,'?') for mid in card.get('idMembers',[])]

    total_c = completed_c = 0
    cl_info = []
    for cid in card.get('idChecklists',[]):
        for cl in trello_data.get('checklists',[]):
            if cl['id'] == cid:
                items = cl.get('checkItems',[])
                total_c     += len(items)
                completed_c += sum(1 for i in items if i.get('state')=='complete')
                cl_info.append(cl.get('name','') + ': ' + str(completed_c) + '/' + str(total_c))

    comments = [
        a.get('date','')[:10] + ': ' + a.get('data',{}).get('text','')
        for a in trello_data.get('actions',[])
        if a.get('type')=='commentCard'
        and a.get('data',{}).get('card',{}).get('id')==card.get('id')
    ]

    task_data.append({
        'name':         nombre,
        'status':       list_dict.get(card.get('idList'),'?'),
        'members':      ', '.join(members) or 'Sin asignar',
        'due_date':     card.get('due',''),
        'labels':       ', '.join(labels),
        'norma':        next((n for n in NORMAS if n in labels), 'General'),
        'prioridad':    next((p for p in PRIORIDADES if p in labels), 'Sin definir'),
        'checklists':   '; '.join(cl_info) or 'Sin checklists',
        'completion':   str(completed_c) + '/' + str(total_c) if total_c else 'N/A',
        'completion_pct': round(completed_c/total_c*100,1) if total_c else 0,
        'comments':     '; '.join(comments) or 'Sin comentarios',
        'url':          card.get('url',''),
    })

tasks_df = pd.DataFrame(task_data)
tasks_df['due_date'] = pd.to_datetime(tasks_df['due_date'], errors='coerce').dt.tz_localize(None)

xlsx_path = os.path.join(SCRIPT_DIR, 'tareas_sgc_' + fecha_str + '.xlsx')
csv_path  = os.path.join(SCRIPT_DIR, 'tareas_sgc_' + fecha_str + '.csv')
tasks_df.to_excel(xlsx_path, index=False)
tasks_df.to_csv(csv_path, index=False)
print("  OK Excel : " + os.path.basename(xlsx_path))
print("  OK CSV   : " + os.path.basename(csv_path) + "\n")

mask_term = tasks_df['status'].str.lower().str.contains('termin|complet|done|finaliz', na=False)
mask_proc = tasks_df['status'].str.lower().str.contains('en proceso|en curso|doing|progres|revision', na=False)
mask_pend = ~mask_term & ~mask_proc

terminadas = tasks_df[mask_term]
en_proceso = tasks_df[mask_proc]
pendientes = tasks_df[mask_pend]
total      = len(tasks_df)
pct_avance = round(len(terminadas)/total*100, 1) if total > 0 else 0

tasks_df['is_overdue'] = (~mask_term) & tasks_df['due_date'].notna() & (tasks_df['due_date'] < pd.Timestamp(hoy))
tasks_df['is_soon']    = (~mask_term) & (~tasks_df['is_overdue']) & tasks_df['due_date'].notna() & \
                         (tasks_df['due_date'] <= pd.Timestamp(hoy + timedelta(days=7)))
overdue_tasks = tasks_df[tasks_df['is_overdue']].sort_values('due_date')
soon_tasks    = tasks_df[tasks_df['is_soon']].sort_values('due_date')

# PARTE 2 - Informe HTML para el equipo
print("[2/3] Generando informe HTML para el equipo...")

def fig_to_b64(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight', dpi=100)
    buf.seek(0)
    b64 = base64.b64encode(buf.read()).decode()
    plt.close(fig)
    return b64

fig, ax = plt.subplots(figsize=(7,4))
sc = tasks_df['status'].value_counts()
ax.bar(sc.index, sc.values, color=['#2196F3','#4CAF50','#FFC107','#9E9E9E'][:len(sc)], edgecolor='white')
ax.set_title('Distribucion por Estado', fontweight='bold'); ax.set_ylabel('N acciones')
plt.xticks(rotation=20, ha='right'); plt.tight_layout()
c1 = fig_to_b64(fig)

fig, ax = plt.subplots(figsize=(5,5))
sizes = [len(terminadas), len(en_proceso), len(pendientes)]
if sum(sizes) > 0:
    ax.pie(sizes, labels=['Terminadas','En Proceso','Pendientes'],
           colors=['#4CAF50','#2196F3','#FFC107'], autopct='%1.0f%%',
           startangle=90, explode=(0.05,0.05,0.05))
ax.set_title('Progreso General', fontweight='bold'); plt.tight_layout()
c2 = fig_to_b64(fig)

fig, ax = plt.subplots(figsize=(8,4))
am = tasks_df[~mask_term]['members'].str.split(', ').explode()
am = am[am != 'Sin asignar'].value_counts().head(10)
if len(am) > 0:
    am.index = [n.split()[0] if len(n)>15 else n for n in am.index]
    ax.bar(am.index, am.values, color='#2196F3', edgecolor='white')
ax.set_title('Carga por Responsable (activas)', fontweight='bold')
ax.set_ylabel('N acciones'); plt.xticks(rotation=25, ha='right'); plt.tight_layout()
c3 = fig_to_b64(fig)

fig, ax = plt.subplots(figsize=(6,4))
nc = tasks_df['norma'].value_counts()
ax.barh(nc.index[::-1], nc.values[::-1], color='#1F3864', edgecolor='white')
ax.set_title('Acciones por Norma/Area', fontweight='bold')
ax.set_xlabel('N acciones'); plt.tight_layout()
c4 = fig_to_b64(fig)

CSS = (
    'body{font-family:Arial,sans-serif;background:#f0f4f8;color:#333;padding:20px}'
    '.hdr{background:linear-gradient(135deg,#1F3864,#2E75B6);color:white;padding:22px 28px;border-radius:10px;margin-bottom:20px}'
    '.hdr h1{font-size:20px;margin:0 0 4px}.hdr p{font-size:13px;opacity:.85}'
    '.metrics{display:flex;flex-wrap:wrap;gap:12px;margin-bottom:20px}'
    '.m{flex:1;min-width:110px;background:white;border-radius:8px;padding:14px;text-align:center;border-top:4px solid #2196F3;box-shadow:0 1px 4px rgba(0,0,0,.08)}'
    '.m.done{border-top-color:#4CAF50}.m.ov{border-top-color:#f44336}.m.warn{border-top-color:#FF9800}'
    '.mv{font-size:26px;font-weight:bold;color:#1F3864}.ml{font-size:11px;color:#777;margin-top:4px}'
    '.charts{display:flex;flex-wrap:wrap;gap:14px;margin-bottom:20px}'
    '.chart{background:white;border-radius:8px;padding:10px;flex:1;min-width:260px;box-shadow:0 1px 4px rgba(0,0,0,.08)}'
    '.chart img{width:100%}'
    'h2{color:#1F3864;font-size:15px;border-bottom:2px solid #2196F3;padding-bottom:5px;margin:20px 0 10px}'
    'table{width:100%;border-collapse:collapse;font-size:12px;background:white;border-radius:8px;overflow:hidden;box-shadow:0 1px 4px rgba(0,0,0,.08);margin-bottom:14px}'
    'th{background:#1F3864;color:white;padding:8px 10px;text-align:left;font-size:11px;text-transform:uppercase}'
    'td{padding:7px 10px;border-bottom:1px solid #eee}'
    'tr:nth-child(even) td{background:#f7fafd}.ov td{background:#ffe8e8!important}'
    '.ba{display:inline-block;padding:2px 8px;border-radius:10px;font-size:10px;color:white;font-weight:bold}'
    '.ba-a{background:#f44336}.ba-m{background:#FF9800}.ba-b{background:#4CAF50}'
    '.alert{background:#ffe8e8;border-left:4px solid #f44336;padding:10px;border-radius:5px;margin-bottom:10px}'
    '.ok{background:#e8f5e9;border-left:4px solid #4CAF50;padding:10px;border-radius:5px;margin-bottom:10px}'
    '.warn2{background:#fff8e1;border-left:4px solid #FF9800;padding:10px;border-radius:5px;margin-bottom:10px}'
    'footer{text-align:center;color:#aaa;font-size:11px;margin-top:28px}'
)

h  = '<!DOCTYPE html><html lang="es"><head><meta charset="UTF-8">'
h += '<title>Informe Interno SGC</title><style>' + CSS + '</style></head><body>'
h += '<div class="hdr"><h1>Informe Interno - ' + LAB_NAME + '</h1>'
h += '<p>Semana N ' + str(semana_num) + ' | ' + hoy.strftime('%d/%m/%Y') + ' | Generado automaticamente desde Trello</p></div>'
h += '<div class="metrics">'
h += '<div class="m"><div class="mv">' + str(total) + '</div><div class="ml">Total</div></div>'
h += '<div class="m done"><div class="mv">' + str(len(terminadas)) + '</div><div class="ml">Terminadas</div></div>'
h += '<div class="m"><div class="mv">' + str(len(en_proceso)) + '</div><div class="ml">En Proceso</div></div>'
h += '<div class="m"><div class="mv">' + str(len(pendientes)) + '</div><div class="ml">Pendientes</div></div>'
h += '<div class="m done"><div class="mv">' + str(pct_avance) + '%</div><div class="ml">Avance</div></div>'
h += '<div class="m ov"><div class="mv">' + str(len(overdue_tasks)) + '</div><div class="ml">Vencidas</div></div>'
h += '<div class="m warn"><div class="mv">' + str(len(soon_tasks)) + '</div><div class="ml">Vencen 7 dias</div></div>'
h += '</div><div class="charts">'
for img in [c1,c2,c3,c4]:
    h += '<div class="chart"><img src="data:image/png;base64,' + img + '"/></div>'
h += '</div><h2>Alertas</h2>'

if len(overdue_tasks) > 0:
    h += '<div class="alert"><b>Atencion: ' + str(len(overdue_tasks)) + ' accion(es) vencida(s)</b></div>'
    h += '<table><tr><th>Accion</th><th>Responsable</th><th>Vencia</th><th>Estado</th><th>Dias retraso</th></tr>'
    for _, r in overdue_tasks.iterrows():
        dias = (pd.Timestamp(hoy) - r['due_date']).days
        due  = r['due_date'].strftime('%d/%m/%Y') if pd.notna(r['due_date']) else '-'
        h += '<tr class="ov"><td>' + r['name'] + '</td><td>' + r['members'] + '</td><td>' + due + '</td><td>' + r['status'] + '</td><td>' + str(dias) + '</td></tr>'
    h += '</table>'
else:
    h += '<div class="ok">No hay acciones vencidas</div>'

if len(soon_tasks) > 0:
    h += '<div class="warn2"><b>' + str(len(soon_tasks)) + ' accion(es) vencen en 7 dias</b></div>'
    h += '<table><tr><th>Accion</th><th>Responsable</th><th>Vence</th><th>Estado</th><th>Dias restantes</th></tr>'
    for _, r in soon_tasks.iterrows():
        dias = (r['due_date'] - pd.Timestamp(hoy)).days
        due  = r['due_date'].strftime('%d/%m/%Y') if pd.notna(r['due_date']) else '-'
        h += '<tr><td>' + r['name'] + '</td><td>' + r['members'] + '</td><td>' + due + '</td><td>' + r['status'] + '</td><td>' + str(dias) + '</td></tr>'
    h += '</table>'

h += '<h2>Estado Detallado</h2>'
h += '<table><tr><th>Accion</th><th>Norma</th><th>Estado</th><th>Responsable</th><th>Vence</th><th>Prioridad</th><th>Checklist</th></tr>'
for _, r in tasks_df.sort_values(['status','due_date']).iterrows():
    due = r['due_date'].strftime('%d/%m/%Y') if pd.notna(r['due_date']) else '-'
    p   = r.get('prioridad','')
    if 'alta'  in p.lower(): badge = '<span class="ba ba-a">Alta</span>'
    elif 'media' in p.lower(): badge = '<span class="ba ba-m">Media</span>'
    elif 'baja' in p.lower(): badge = '<span class="ba ba-b">Baja</span>'
    else: badge = '-'
    cls = ' class="ov"' if r['is_overdue'] else ''
    h += '<tr' + cls + '><td>' + r['name'] + '</td><td>' + r['norma'] + '</td><td>' + r['status'] + '</td><td>' + r['members'] + '</td><td>' + due + '</td><td>' + badge + '</td><td>' + r['completion'] + '</td></tr>'
h += '</table>'
h += '<footer>Generado el ' + datetime.now().strftime('%d/%m/%Y %H:%M') + ' - ' + LAB_NAME + '</footer></body></html>'

html_path = os.path.join(SCRIPT_DIR, 'informe_interno_' + fecha_str + '.html')
with open(html_path, 'w', encoding='utf-8') as f:
    f.write(h)
print("  OK HTML  : " + os.path.basename(html_path) + "\n")

# PARTE 3 - Informe Word para la Direccion
print("[3/3] Generando informe Word para la Direccion...")

def cell_bg(cell, hex_c):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'), hex_c)
    tcPr.append(shd)

def cell_borders(cell, color='CCCCCC', sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    tb   = OxmlElement('w:tcBorders')
    for e in ('top','left','bottom','right'):
        b = OxmlElement('w:' + e)
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),str(sz))
        b.set(qn('w:space'),'0');   b.set(qn('w:color'), color)
        tb.append(b)
    tcPr.append(tb)

def para_border_bottom(para, color='1F3864', sz=18):
    pPr = para._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    b   = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),str(sz))
    b.set(qn('w:space'),'1');   b.set(qn('w:color'), color)
    pb.append(b); pPr.append(pb)

def sec_title(doc, n, txt):
    p = doc.add_paragraph()
    r = p.add_run(str(n) + '. ' + txt.upper())
    r.font.size = Pt(12); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F,0x36,0x64)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(5)
    para_border_bottom(p, C_AZUL_OSC, 8)

def make_tbl(doc, headers, widths, hdr_bg=C_AZUL_OSC):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    row = t.rows[0]
    for i,(hdr,w) in enumerate(zip(headers,widths)):
        c = row.cells[i]; c.width = Cm(w); c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr); r.bold = True; r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        cell_bg(c, hdr_bg); cell_borders(c, '4472C4', 6)
    return t

def add_tbl_row(t, vals, widths, bg=C_BLANCO, fs=8.5, aligns=None, bc='CCCCCC'):
    row = t.add_row()
    for i,(v,w) in enumerate(zip(vals,widths)):
        c = row.cells[i]; c.width = Cm(w); c.text = ''
        p = c.paragraphs[0]
        if aligns: p.alignment = aligns[i]
        r = p.add_run(str(v) if v is not None else '-')
        r.font.size = Pt(fs)
        cell_bg(c, bg); cell_borders(c, bc, 4)

def fig_to_buf(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight', dpi=130)
    buf.seek(0); plt.close(fig)
    return buf

fig, ax = plt.subplots(figsize=(4.5,4.5))
sz = [len(terminadas), len(en_proceso), len(pendientes)]
if sum(sz) > 0:
    wedges, _, auts = ax.pie(sz, labels=['Terminadas','En Proceso','Pendientes'],
        colors=['#4CAF50','#2196F3','#FFC107'], autopct='%1.0f%%',
        startangle=90, pctdistance=0.78,
        wedgeprops=dict(edgecolor='white', linewidth=1.5))
    for t2 in auts: t2.set_fontsize(12); t2.set_fontweight('bold')
ax.set_title('Avance General - ' + str(pct_avance) + '%', fontsize=11, fontweight='bold', pad=8)
plt.tight_layout()
buf_pie = fig_to_buf(fig)

normas_u = sorted(tasks_df['norma'].unique())
fig, ax = plt.subplots(figsize=(5.5,3.2))
tc = [tasks_df[(tasks_df['norma']==n)&mask_term].shape[0] for n in normas_u]
pc = [tasks_df[(tasks_df['norma']==n)&mask_proc].shape[0] for n in normas_u]
qc = [tasks_df[(tasks_df['norma']==n)&mask_pend].shape[0] for n in normas_u]
x  = range(len(normas_u)); bw = 0.25
ax.bar([i-bw for i in x], tc, width=bw, label='Terminadas', color='#4CAF50')
ax.bar(list(x),            pc, width=bw, label='En Proceso',  color='#2196F3')
ax.bar([i+bw for i in x], qc, width=bw, label='Pendientes',  color='#FFC107')
ax.set_xticks(list(x)); ax.set_xticklabels(normas_u)
ax.set_title('Estado por Norma/Area', fontsize=10, fontweight='bold', pad=8)
ax.legend(fontsize=8); ax.set_ylabel('N Acciones')
ax.yaxis.get_major_locator().set_params(integer=True)
plt.tight_layout()
buf_bar = fig_to_buf(fig)

doc = Document()
s = doc.sections[0]
s.page_width = Cm(21); s.page_height = Cm(29.7)
s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
s.top_margin = Cm(2.2); s.bottom_margin = Cm(2)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

hp = doc.sections[0].header.paragraphs[0]
hp.clear()
hr = hp.add_run(LAB_NAME + ' | Sistema de Gestion de Calidad | Documento de uso interno')
hr.font.size = Pt(7.5); hr.font.color.rgb = RGBColor(0x80,0x80,0x80)
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

fp = doc.sections[0].footer.paragraphs[0]
fp.clear()
fr = fp.add_run('Generado automaticamente el ' + datetime.now().strftime('%d/%m/%Y %H:%M') + ' | ' + LAB_NAME)
fr.font.size = Pt(7.5); fr.font.color.rgb = RGBColor(0x80,0x80,0x80)
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph(); para_border_bottom(p, C_AZUL_MED, 30)
p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph()
r = p.add_run(LAB_NAME.upper())
r.font.size = Pt(18); r.font.bold = True; r.font.color.rgb = RGBColor(0x1F,0x36,0x64)
p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run('INFORME SEMANAL DE SEGUIMIENTO DEL PLAN DE MEJORA')
r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = RGBColor(0x2E,0x75,0xB6)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run('Normas de referencia: ' + '  |  '.join(NORMAS))
r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = RGBColor(0x60,0x60,0x60)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph(); para_border_bottom(p, C_AZUL_MED, 14)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run('Fecha de emision:  ' + fecha_larga + '      Semana N ' + str(semana_num))
r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0x40,0x40,0x40)
p.paragraph_format.space_after = Pt(12)

sec_title(doc, 1, 'Resumen Ejecutivo')

metricas = [
    ('Total de acciones del plan de mejora',       str(total)),
    ('Acciones terminadas',                        str(len(terminadas)) + ' (' + str(pct_avance) + '% del total)'),
    ('Acciones en proceso',                        str(len(en_proceso))),
    ('Acciones pendientes de iniciar',             str(len(pendientes))),
    ('Acciones vencidas (fecha limite superada)',  'ATENCION: ' + str(len(overdue_tasks)) if len(overdue_tasks) else '0'),
    ('Acciones que vencen en los proximos 7 dias', str(len(soon_tasks))),
    ('Responsables involucrados',                  str(tasks_df['members'].str.split(', ').explode().nunique())),
]

tm = make_tbl(doc, ['Indicador','Valor'], [11.5, 4.5])
for i,(ind,val) in enumerate(metricas):
    bg = 'EBF3FB' if i%2==0 else C_BLANCO
    if 'vencidas' in ind and len(overdue_tasks)>0: bg = C_ROJO_CLR
    if 'terminadas' in ind: bg = C_VERDE_CLR
    row = tm.add_row()
    c0 = row.cells[0]; c0.width = Cm(11.5); c0.text = ''
    r0 = c0.paragraphs[0].add_run(ind); r0.font.size = Pt(9)
    cell_bg(c0, bg); cell_borders(c0)
    c1 = row.cells[1]; c1.width = Cm(4.5); c1.text = ''
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = c1.paragraphs[0].add_run(val); r1.font.size = Pt(9); r1.bold = True
    if 'vencidas' in ind and len(overdue_tasks)>0:
        r1.font.color.rgb = RGBColor(0xC0,0x00,0x00)
    cell_bg(c1, bg); cell_borders(c1)

doc.add_paragraph()

tc2 = doc.add_table(rows=1, cols=2); tc2.alignment = WD_TABLE_ALIGNMENT.CENTER
cl = tc2.rows[0].cells[0]; cr = tc2.rows[0].cells[1]
cl.width = Cm(8); cr.width = Cm(8)
cl.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
cr.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
cl.paragraphs[0].add_run().add_picture(buf_pie, width=Cm(7.5))
cr.paragraphs[0].add_run().add_picture(buf_bar, width=Cm(7.5))
for c in [cl, cr]:
    tcPr = c._tc.get_or_add_tcPr()
    tb = OxmlElement('w:tcBorders')
    for e in ('top','left','bottom','right'):
        b = OxmlElement('w:' + e); b.set(qn('w:val'),'none'); tb.append(b)
    tcPr.append(tb)
doc.add_paragraph()

sec_title(doc, 2, 'Distribucion por Norma / Area')
cw2 = [4.0, 2.0, 3.0, 3.5, 3.5]
tn2 = make_tbl(doc, ['Norma/Area','Total','Terminadas','En Proceso','Pendientes'], cw2, C_AZUL_MED)
al2 = [WD_ALIGN_PARAGRAPH.LEFT]+[WD_ALIGN_PARAGRAPH.CENTER]*4
for i,norma in enumerate(sorted(tasks_df['norma'].unique())):
    mn = tasks_df['norma'] == norma
    bg = 'EBF3FB' if i%2==0 else C_BLANCO
    add_tbl_row(tn2, [norma, mn.sum(), (mn&mask_term).sum(), (mn&mask_proc).sum(), (mn&~mask_term&~mask_proc).sum()],
                cw2, bg=bg, aligns=al2, fs=9)
doc.add_paragraph()

sec_title(doc, 3, 'Alertas y Seguimiento')
if len(overdue_tasks)==0 and len(soon_tasks)==0:
    p = doc.add_paragraph()
    r = p.add_run('No se registran acciones vencidas ni proximas a vencer. El plan se desarrolla dentro de los plazos previstos.')
    r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0x00,0x70,0x00)
else:
    if len(overdue_tasks) > 0:
        p = doc.add_paragraph()
        r = p.add_run('Acciones con fecha limite vencida (' + str(len(overdue_tasks)) + '):')
        r.font.size = Pt(10); r.bold = True; r.font.color.rgb = RGBColor(0xC0,0x00,0x00)
        p.paragraph_format.space_after = Pt(4)
        cw = [6.5,3.5,2.5,3.5]
        t3 = make_tbl(doc, ['Accion','Responsable','Fecha Limite','Dias retraso'], cw, C_ROJO)
        for _,r2 in overdue_tasks.iterrows():
            dias = (pd.Timestamp(hoy)-r2['due_date']).days
            due  = r2['due_date'].strftime('%d/%m/%Y') if pd.notna(r2['due_date']) else '-'
            add_tbl_row(t3, [r2['name'],r2['members'],due,str(dias)+' dias'], cw, bg=C_ROJO_CLR, bc='FF9999')
        doc.add_paragraph()
    if len(soon_tasks) > 0:
        p = doc.add_paragraph()
        r = p.add_run('Acciones que vencen en los proximos 7 dias (' + str(len(soon_tasks)) + '):')
        r.font.size = Pt(10); r.bold = True; r.font.color.rgb = RGBColor(0xBB,0x77,0x00)
        p.paragraph_format.space_after = Pt(4)
        cw = [6.5,3.5,2.5,3.5]
        t4 = make_tbl(doc, ['Accion','Responsable','Fecha Limite','Dias restantes'], cw, C_NARANJA)
        for _,r2 in soon_tasks.iterrows():
            dias = (r2['due_date']-pd.Timestamp(hoy)).days
            due  = r2['due_date'].strftime('%d/%m/%Y') if pd.notna(r2['due_date']) else '-'
            add_tbl_row(t4, [r2['name'],r2['members'],due,str(dias)+' dias'], cw, bg=C_NARJ_CLR, bc='FFCC66')
        doc.add_paragraph()

sec_title(doc, 4, 'Estado Detallado de Acciones')
activas = tasks_df[~mask_term].sort_values(['norma','due_date'])
if len(activas) > 0:
    p = doc.add_paragraph()
    r = p.add_run('Acciones en curso y pendientes (' + str(len(activas)) + '):')
    r.font.size = Pt(9.5); r.bold = True; p.paragraph_format.space_after = Pt(4)
    cwd = [5.5,2.0,3.5,2.5,2.5]
    t5  = make_tbl(doc, ['Accion','Norma','Responsable','Fecha Limite','Estado'], cwd)
    ald = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
           WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
    for i,(_,r2) in enumerate(activas.iterrows()):
        due = r2['due_date'].strftime('%d/%m/%Y') if pd.notna(r2['due_date']) else '-'
        bg  = C_ROJO_CLR if r2['is_overdue'] else ('EBF3FB' if i%2==0 else C_BLANCO)
        add_tbl_row(t5, [r2['name'],r2['norma'],r2['members'],due,r2['status']], cwd, bg=bg, aligns=ald)
    doc.add_paragraph()
if len(terminadas) > 0:
    p = doc.add_paragraph()
    r = p.add_run('Acciones terminadas (' + str(len(terminadas)) + '):')
    r.font.size = Pt(9.5); r.bold = True; r.font.color.rgb = RGBColor(0x00,0x70,0x00)
    p.paragraph_format.space_after = Pt(4)
    cwt = [7.5,2.0,6.5]
    t6  = make_tbl(doc, ['Accion','Norma','Responsable'], cwt, C_VERDE)
    for i,(_,r2) in enumerate(terminadas.iterrows()):
        add_tbl_row(t6, [r2['name'],r2['norma'],r2['members']],
                    cwt, bg=C_VERDE_CLR if i%2==0 else C_BLANCO, bc='99CC99')
    doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.page_break_before = True
p.paragraph_format.space_after = Pt(0)
sec_title(doc, 5, 'Observaciones y Compromisos')

if observaciones:
    fuente = observaciones.get('fuente_reuniones', [])
    if fuente:
        p = doc.add_paragraph()
        r = p.add_run('Basado en las reuniones de seguimiento: ' + ', '.join(fuente))
        r.font.size = Pt(8.5); r.font.italic = True; r.font.color.rgb = RGBColor(0x80,0x80,0x80)
        p.paragraph_format.space_after = Pt(10)
else:
    p = doc.add_paragraph()
    r = p.add_run('Completar en reunion de seguimiento. Incluir novedades, decisiones y compromisos para la proxima semana.')
    r.font.size = Pt(8.5); r.font.italic = True; r.font.color.rgb = RGBColor(0x80,0x80,0x80)
    p.paragraph_format.space_after = Pt(10)

campos = [
    ('novedades',   'Novedades del periodo:'),
    ('decisiones',  'Decisiones adoptadas:'),
    ('compromisos', 'Compromisos para la proxima semana:'),
]
for key, campo in campos:
    p = doc.add_paragraph()
    r = p.add_run(campo); r.font.size = Pt(10); r.bold = True
    r.font.color.rgb = RGBColor(0x2E,0x75,0xB6)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)

    items = observaciones.get(key) if observaciones else None
    if items:
        for it in items:
            p = doc.add_paragraph(style='List Bullet')
            r = p.add_run(it); r.font.size = Pt(9.5)
            p.paragraph_format.space_after = Pt(3)
    else:
        for _ in range(4):
            p = doc.add_paragraph(); para_border_bottom(p,'AAAAAA',4)
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(5)

p = doc.add_paragraph()
r = p.add_run('Firma Responsable SGC:                               Firma Direccion del Laboratorio:')
r.font.size = Pt(10); r.bold = True
r.font.color.rgb = RGBColor(0x2E,0x75,0xB6)
p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
for _ in range(3):
    p = doc.add_paragraph(); para_border_bottom(p,'AAAAAA',4)
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(5)

docx_path = os.path.join(SCRIPT_DIR, 'informe_direccion_' + fecha_str + '.docx')
doc.save(docx_path)
print("  OK Word  : " + os.path.basename(docx_path) + "\n")

print(SEP)
print('  Proceso completado! Archivos generados:')
print('    - ' + os.path.basename(xlsx_path))
print('    - ' + os.path.basename(csv_path))
print('    - ' + os.path.basename(html_path))
print('    - ' + os.path.basename(docx_path))
print(SEP + '\n')
